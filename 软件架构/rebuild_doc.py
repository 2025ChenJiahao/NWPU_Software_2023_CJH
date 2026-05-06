# -*- coding: utf-8 -*-
"""重建建模作业Word文档 - 使用SVG图片，优化排版"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import cairosvg
import os
import io

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SVG_DIR = os.path.join(BASE_DIR, 'svg')
OUTPUT = os.path.join(BASE_DIR, '共享自习室预约管理系统_建模作业.docx')

# SVG -> PNG conversion
def svg_to_png(svg_path, scale=2.0):
    png_data = cairosvg.svg2png(url=svg_path, scale=scale)
    return io.BytesIO(png_data)

doc = Document()

# ========== Global styles ==========
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    hs.paragraph_format.space_after = Pt(12 if level == 1 else 8)
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(15)
    else:
        hs.font.size = Pt(13)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

def add_title(text, level=1):
    return doc.add_heading(text, level=level)

def add_para(text, bold=False, indent=False, font_size=None, alignment=None, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return p

def add_svg_image(svg_filename, caption="", width_inches=5.5):
    svg_path = os.path.join(SVG_DIR, svg_filename)
    if not os.path.exists(svg_path):
        add_para(f'[图片缺失: {svg_filename}]', color=RGBColor(255, 0, 0))
        return
    png_stream = svg_to_png(svg_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_stream, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(100, 100, 100)
        r.italic = True

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        from docx.oxml import OxmlElement
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '003366')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            if r_idx % 2 == 0:
                from docx.oxml import OxmlElement
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F0F4F8')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table

# ==================== COVER ====================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('软件架构与设计模式')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('建模作业')
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 40)
run.font.color.rgb = RGBColor(0, 120, 215)
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('共享自习室预约管理系统')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 82, 136)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UML面向对象建模')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(4):
    doc.add_paragraph()

info = [
    ('课程名称', '软件架构与设计模式'),
    ('作业题目', '共享自习室预约管理系统 UML建模'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(13)

doc.add_page_break()

# ==================== TOC ====================
add_title('目录', 1)
toc_items = [
    ('一、系统概述', 0),
    ('二、用例建模', 0),
    ('2.1 参与者识别', 1),
    ('2.2 用例图', 1),
    ('2.3 用例详细说明', 1),
    ('三、静态建模', 0),
    ('3.1 上下文类图', 1),
    ('3.2 实体类设计', 1),
    ('3.3 边界类设计', 1),
    ('3.4 控制类设计', 1),
    ('3.5 应用逻辑类设计', 1),
    ('3.6 类关系总结', 1),
    ('3.7 用例-类映射', 1),
    ('四、动态建模', 0),
    ('4.1 通信图：预约座位', 1),
    ('4.2 消息序列说明', 1),
    ('4.3 状态图：预约状态', 1),
    ('五、总结', 0),
]
for item, level in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(('    ' * level) + item)
    run.font.size = Pt(12 if level == 0 else 11)
    run.bold = (level == 0)

doc.add_page_break()

# ==================== 1. 系统概述 ====================
add_title('一、系统概述', 1)

add_para('本系统为高校及城市共享自习室提供在线预约与运营管理平台，支持用户按时段预约座位、在线支付、扫码签到入场，以及管理员对自习室和订单的统一管理。', indent=True)

add_para('系统核心功能包括：', bold=True)
features = [
    '用户注册与登录（对接外部用户账户系统）',
    '自习室浏览与座位时段查询',
    '在线预约座位（一时段一座一用户约束）',
    '预约订单生成与在线支付（对接外部支付系统）',
    '扫码签到与入场验证（对接外部门禁控制系统）',
    '预约管理（查看、取消、历史记录）',
    '信用管理（爽约记录、连续爽约限制）',
    '管理员后台（自习室管理、订单查看、异常标记）',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

add_para('系统边界与外部依赖：', bold=True)
add_para('本系统依赖三个外部系统：用户账户系统（负责注册、登录、身份认证）、支付系统（负责在线支付与退款）、门禁控制系统（负责扫码验证与入场控制）。系统通过标准接口与外部系统交互，自身不存储用户密码等敏感认证信息。', indent=True)

doc.add_page_break()

# ==================== 2. 用例建模 ====================
add_title('二、用例建模', 1)

# 2.1
add_title('2.1 参与者识别', 2)
add_para('根据系统需求分析，识别出以下参与者（Actor）：')
add_table(
    ['参与者', '类型', '描述'],
    [
        ['普通用户', '主要参与者', '使用自习室进行自习的人，可浏览、预约、支付、签到、管理预约'],
        ['管理员', '主要参与者', '运营管理人员，负责自习室信息维护、订单监控、异常处理'],
        ['用户账户系统', '外部系统', '提供用户注册、登录、身份认证服务'],
        ['支付系统', '外部系统', '提供在线支付、退款等金融服务'],
        ['门禁控制系统', '外部系统', '提供扫码验证、开门控制等入场服务'],
    ]
)

# 2.2
add_title('2.2 用例图', 2)
add_para('以下为系统的用例图，展示了系统中所有参与者及其对应的用例关系：')
add_svg_image('1.svg', '图2-1 共享自习室预约管理系统用例图', width_inches=5.0)

add_para('用例清单：', bold=True)
add_table(
    ['编号', '用例名称', '参与者', '简要描述'],
    [
        ['UC01', '注册/登录', '普通用户、用户账户系统', '用户通过账户系统完成注册或登录认证'],
        ['UC02', '浏览自习室', '普通用户', '查看可用自习室列表及基本信息'],
        ['UC03', '查看时段座位状态', '普通用户', '查看指定自习室某时段的座位占用情况'],
        ['UC04', '预约座位', '普通用户', '选择时段和座位提交预约'],
        ['UC05', '支付订单', '普通用户、支付系统', '对预约订单进行在线支付'],
        ['UC06', '扫码签到', '普通用户、门禁控制系统', '预约时段内扫码验证入场'],
        ['UC07', '查看预约记录', '普通用户', '查看当前和历史预约记录'],
        ['UC08', '取消预约', '普通用户', '在预约开始前取消预约'],
        ['UC09', '管理自习室信息', '管理员', '对自习室信息进行增删改操作'],
        ['UC10', '查看所有订单', '管理员', '查看系统中所有预约订单'],
        ['UC11', '标记异常', '管理员', '标记未使用、违规占座等异常订单'],
    ]
)

# 2.3
add_title('2.3 用例详细说明', 2)

add_title('用例1：UC04 预约座位', 3)
add_table(
    ['项目', '内容'],
    [
        ['用例名称', '预约座位'],
        ['用例编号', 'UC04'],
        ['参与者', '普通用户（已登录）'],
        ['前置条件', '1. 用户已登录系统\n2. 系统中存在可用自习室和空闲座位'],
        ['后置条件', '1. 系统生成一条"待支付"状态的预约订单\n2. 对应座位在该时段被锁定，其他用户不可预约'],
        ['基本事件流', '1. 用户选择目标自习室\n2. 用户选择预约日期和时段\n3. 系统显示该时段座位状态图（空闲/已占/锁定）\n4. 用户选择一个空闲座位\n5. 系统校验：该用户在该时段无其他预约\n6. 系统锁定座位（暂态，5分钟有效）\n7. 系统生成预约订单（状态：待支付）\n8. 系统跳转至支付流程（UC05）'],
        ['备选事件流', '3a. 该时段所有座位已满 → 系统提示"该时段已满，请选择其他时段"，返回步骤2\n5a. 用户在该时段已有预约 → 系统提示"您在该时段已有预约"，返回步骤2\n6a. 座位锁定失败（被其他用户抢先选择） → 系统提示"该座位已被选择"，返回步骤3'],
        ['异常事件流', '7a. 系统生成订单失败 → 系统提示"预约失败，请重试"，释放座位锁定\n6b. 锁定超时（5分钟未支付） → 系统自动取消订单，释放座位'],
        ['业务规则', 'BR1: 一时段一座一用户，同一用户同一时段只能预约一个座位\nBR2: 座位锁定有效期为5分钟\nBR3: 可预约多个不同时段\nBR4: 预约时间不得早于当前时间'],
    ]
)

add_title('用例2：UC05 预约确认与支付', 3)
add_table(
    ['项目', '内容'],
    [
        ['用例名称', '预约确认与支付'],
        ['用例编号', 'UC05'],
        ['参与者', '普通用户、支付系统'],
        ['前置条件', '1. 用户已登录系统\n2. 已生成一条"待支付"状态的预约订单\n3. 座位已锁定（5分钟有效期内）'],
        ['后置条件', '支付成功：订单状态变为"已预约"，座位确认分配\n支付失败/超时：订单取消，座位释放'],
        ['基本事件流', '1. 系统展示订单详情（自习室、时段、座位、金额）\n2. 用户确认订单信息\n3. 用户选择支付方式\n4. 系统调用支付系统创建支付单\n5. 支付系统返回支付页面/二维码\n6. 用户完成支付\n7. 支付系统回调通知支付结果\n8. 系统更新订单状态为"已预约"\n9. 系统向用户展示预约成功页面（含入场二维码）'],
        ['备选事件流', '2a. 用户取消支付 → 系统取消订单，释放座位，返回预约页面\n6a. 支付失败（余额不足、网络异常等） → 支付系统返回失败原因，系统提示用户重试或更换支付方式\n6b. 用户在支付页面超时 → 系统自动取消订单，释放座位'],
        ['异常事件流', '4a. 支付系统不可用 → 系统提示"支付服务暂时不可用，请稍后重试"，保留订单5分钟\n7a. 支付回调通知异常（已支付但回调失败） → 系统设置定时对账任务，30分钟内未收到成功回调则主动查询支付系统确认状态'],
        ['业务规则', 'BR5: 支付限时5分钟，超时自动取消\nBR6: 支付成功后不可退款（预约开始前可取消预约，费用原路退回）\nBR7: 系统记录支付流水号，用于对账和退款'],
    ]
)

doc.add_page_break()

# ==================== 3. 静态建模 ====================
add_title('三、静态建模', 1)

# 3.1
add_title('3.1 上下文类图', 2)
add_para('上下文类图展示系统与外部系统的交互接口，明确系统边界。')
add_svg_image('2.svg', '图3-1 上下文类图', width_inches=4.0)

# 3.2
add_title('3.2 实体类设计', 2)
add_para('实体类对应系统中的核心业务数据对象，通常映射为数据库表。')
add_table(
    ['类名', '属性', '类型', '说明'],
    [
        ['User\n(用户)', 'userId : String\nname : String\nphone : String\nemail : String\ncreditScore : Integer\ncreateTime : DateTime', '实体类', '系统注册用户，存储基本信息和信用分'],
        ['StudyRoom\n(自习室)', 'roomId : String\nname : String\nlocation : String\ncapacity : Integer\nopenTime : Time\ncloseTime : Time\nstatus : RoomStatus', '实体类', '自习室基本信息，含容量和营业时间'],
        ['Seat\n(座位)', 'seatId : String\nroomId : String\nseatNo : String\nseatType : SeatType\nstatus : SeatStatus', '实体类', '自习室中的座位，关联所属自习室'],
        ['TimeSlot\n(时段)', 'slotId : String\nroomId : String\nstartTime : DateTime\nendTime : DateTime\nprice : Decimal', '实体类', '可预约的时间段，含价格信息'],
        ['Reservation\n(预约)', 'reservationId : String\nuserId : String\nslotId : String\nseatId : String\nstatus : ReservationStatus\ncreateTime : DateTime\ncheckinTime : DateTime', '实体类', '预约记录，核心业务实体'],
        ['Order\n(订单)', 'orderId : String\nreservationId : String\nuserId : String\namount : Decimal\nstatus : OrderStatus\npayTime : DateTime\npayTransId : String', '实体类', '支付订单，关联预约记录'],
        ['Violation\n(异常记录)', 'violationId : String\nreservationId : String\ntype : ViolationType\ndescription : String\ncreateTime : DateTime\nhandledBy : String', '实体类', '管理员标记的异常记录'],
    ]
)

add_para('关联枚举类型：', bold=True)
add_table(
    ['枚举名', '取值', '说明'],
    [
        ['ReservationStatus', 'PENDING_PAYMENT, CONFIRMED, IN_USE, COMPLETED, CANCELLED, NO_SHOW, EXPIRED', '预约状态'],
        ['OrderStatus', 'PENDING, PAID, REFUNDED, CANCELLED, EXPIRED', '订单状态'],
        ['SeatStatus', 'AVAILABLE, OCCUPIED, LOCKED, MAINTENANCE', '座位状态'],
        ['SeatType', 'NORMAL, QUIET, POWER, GROUP', '座位类型（普通/静音/带电源/小组）'],
        ['RoomStatus', 'OPEN, CLOSED, MAINTENANCE', '自习室状态'],
        ['ViolationType', 'NO_SHOW, OCCUPY_SEAT, OTHER', '异常类型'],
    ]
)

# 3.3
add_title('3.3 边界类设计', 2)
add_para('边界类对应系统的用户界面或外部接口，负责与参与者交互。')
add_table(
    ['类名', '类型', '对应参与者', '职责'],
    [
        ['LoginPage', 'UI边界', '普通用户', '用户登录/注册界面，调用用户账户系统认证'],
        ['RoomBrowsePage', 'UI边界', '普通用户', '自习室浏览界面，展示自习室列表和详情'],
        ['SeatMapPage', 'UI边界', '普通用户', '座位状态图界面，可视化展示座位占用情况'],
        ['ReservationPage', 'UI边界', '普通用户', '预约操作界面，选择时段和座位'],
        ['PaymentPage', 'UI边界', '普通用户', '支付界面，展示订单详情和支付方式选择'],
        ['CheckInPage', 'UI边界', '普通用户', '签到界面，展示/扫描入场二维码'],
        ['MyReservationPage', 'UI边界', '普通用户', '我的预约界面，查看和管理预约记录'],
        ['AdminRoomPage', 'UI边界', '管理员', '自习室管理界面，增删改查自习室信息'],
        ['AdminOrderPage', 'UI边界', '管理员', '订单管理界面，查看所有订单和标记异常'],
        ['AuthSystemGateway', '系统边界', '用户账户系统', '用户账户系统接口适配器'],
        ['PaymentGateway', '系统边界', '支付系统', '支付系统接口适配器'],
        ['AccessControlGateway', '系统边界', '门禁控制系统', '门禁系统接口适配器'],
    ]
)

# 3.4
add_title('3.4 控制类设计', 2)
add_para('控制类封装用例的业务逻辑，协调边界类和实体类的交互。')
add_table(
    ['类名', '对应用例', '职责'],
    [
        ['AuthController', 'UC01 注册/登录', '处理用户认证流程，调用AuthSystemGateway完成登录/注册'],
        ['RoomController', 'UC02 浏览自习室', '查询自习室列表，按条件筛选'],
        ['SeatController', 'UC03 查看座位状态', '查询指定时段座位占用状态，生成座位图数据'],
        ['ReservationController', 'UC04 预约座位', '处理预约逻辑：校验约束、锁定座位、创建预约、触发支付'],
        ['PaymentController', 'UC05 支付订单', '处理支付流程：创建支付单、回调处理、状态同步'],
        ['CheckInController', 'UC06 扫码签到', '处理签到逻辑：验证二维码、记录签到、调用门禁'],
        ['RecordController', 'UC07 查看预约记录', '查询用户预约历史，按状态/时间筛选'],
        ['CancelController', 'UC08 取消预约', '处理取消逻辑：校验取消条件、释放座位、触发退款'],
        ['RoomManageController', 'UC09 管理自习室', '处理自习室增删改操作'],
        ['OrderManageController', 'UC10/UC11 订单管理', '查询订单列表，标记异常'],
    ]
)

# 3.5
add_title('3.5 应用逻辑类设计', 2)
add_para('应用逻辑类封装跨用例的通用业务规则和算法。')
add_table(
    ['类名', '职责', '关键方法'],
    [
        ['CreditManager', '信用分管理：扣分、恢复、查询', 'deductCredit(userId, score)\nrestoreCredit(userId, score)\ngetCreditLevel(userId)'],
        ['SchedulerService', '定时任务：超时取消、对账、爽约统计', 'autoCancelExpiredReservations()\nreconcilePayments()\nmarkNoShowReservations()'],
        ['NotificationService', '消息通知：预约成功、支付提醒、爽约警告', 'sendReservationConfirm(userId, reservation)\nsendPaymentReminder(userId, order)\nsendNoShowWarning(userId)'],
        ['PricingCalculator', '价格计算：按时段计费、优惠折扣', 'calculatePrice(slotId, userId)\napplyDiscount(amount, couponId)'],
        ['AvailabilityChecker', '可用性检查：座位冲突、时段校验', 'checkSeatAvailable(seatId, slotId)\ncheckUserConflict(userId, slotId)\nvalidateTimeSlot(slotId)'],
    ]
)

# 3.6
add_title('3.6 类关系总结', 2)
add_svg_image('3.svg', '图3-2 类关系图', width_inches=5.5)

add_para('关系说明：', bold=True)
add_table(
    ['关系', '类型', '多重性', '说明'],
    [
        ['User → Reservation', '关联', '1..*', '一个用户可创建多条预约'],
        ['StudyRoom → Seat', '组合', '1..*', '自习室包含多个座位，座位依附于自习室'],
        ['StudyRoom → TimeSlot', '组合', '1..*', '自习室提供多个可预约时段'],
        ['Reservation → Seat', '关联', '*..1', '每条预约对应一个座位'],
        ['Reservation → TimeSlot', '关联', '*..1', '每条预约对应一个时段'],
        ['Reservation → Order', '关联', '1..0..1', '每条预约关联至多一个订单'],
        ['Reservation → Violation', '关联', '1..*', '预约可能产生多条异常记录'],
        ['User → Order', '关联', '1..*', '一个用户持有多条订单'],
    ]
)

# 3.7
add_title('3.7 用例-类映射', 2)
add_para('以下表格展示每个用例涉及的边界类、控制类和实体类：')
add_table(
    ['用例', '边界类', '控制类', '实体类/应用逻辑类'],
    [
        ['UC04 预约座位', 'SeatMapPage\nReservationPage', 'ReservationController', 'User, Seat, TimeSlot, Reservation\nAvailabilityChecker'],
        ['UC05 支付订单', 'PaymentPage\nPaymentGateway', 'PaymentController', 'Order, Reservation\nPricingCalculator'],
        ['UC06 扫码签到', 'CheckInPage\nAccessControlGateway', 'CheckInController', 'Reservation\nCreditManager'],
        ['UC08 取消预约', 'MyReservationPage', 'CancelController', 'Reservation, Order\nCreditManager'],
        ['UC11 标记异常', 'AdminOrderPage', 'OrderManageController', 'Reservation, Violation'],
    ]
)

doc.add_page_break()

# ==================== 4. 动态建模 ====================
add_title('四、动态建模', 1)

# 4.1
add_title('4.1 通信图：预约座位（UC04）', 2)
add_para('选择"预约座位"用例绘制通信图，展示对象之间的消息交互顺序。')
add_svg_image('4.svg', '图4-1 预约座位通信图', width_inches=6.0)

# 4.2
add_title('4.2 消息序列说明', 2)

add_para('正常流程消息序列：', bold=True)
add_table(
    ['序号', '发送方 → 接收方', '消息', '说明'],
    [
        ['1', '用户 → SeatMapPage', '选择自习室和时段', '用户输入查询条件'],
        ['2', 'SeatMapPage → ReservationController', 'requestSeatMap(roomId, slotId)', '请求座位状态数据'],
        ['3', 'ReservationController → AvailabilityChecker', 'checkAvailability(roomId, slotId)', '校验时段有效性并查询可用座位'],
        ['4', 'AvailabilityChecker → Seat', 'querySeats(roomId, slotId)', '查询数据库中座位状态'],
        ['5', 'Seat → AvailabilityChecker', '返回座位状态列表', '包含每个座位的占用/空闲/锁定状态'],
        ['6', 'AvailabilityChecker → ReservationController', '返回可用座位', '过滤后的空闲座位列表'],
        ['7', 'ReservationController → SeatMapPage', '展示座位状态图', '可视化数据返回前端'],
        ['8', '用户 → SeatMapPage', '选择座位', '用户点击目标空闲座位'],
        ['9', 'SeatMapPage → ReservationController', 'reserveSeat(userId, seatId, slotId)', '提交预约请求'],
        ['10', 'ReservationController → AvailabilityChecker', 'checkUserConflict(userId, slotId)', '校验用户在该时段是否已有预约'],
        ['11', 'AvailabilityChecker → ReservationController', '无冲突', '校验通过'],
        ['12', 'ReservationController → Seat', 'lockSeat(seatId, 5min)', '临时锁定座位，防止并发冲突'],
        ['13', 'Seat → ReservationController', '锁定成功', '座位状态变为LOCKED'],
        ['14', 'ReservationController → Reservation', 'createReservation(userId, seatId, slotId)', '创建预约记录（状态：待支付）'],
        ['15', 'Reservation → ReservationController', '返回reservationId', '预约创建成功'],
        ['16', 'ReservationController → Order', 'createOrder(reservationId, amount)', '生成支付订单'],
        ['17', 'Order → ReservationController', '返回orderId', '订单创建成功'],
        ['18', 'ReservationController → SeatMapPage', '返回订单信息', '跳转支付流程'],
        ['19', 'SeatMapPage → PaymentController', 'initiatePayment(orderId)', '发起支付请求'],
        ['20', 'PaymentController → PaymentGateway', 'createPayRequest(orderId, amount)', '调用外部支付系统'],
        ['21', 'PaymentGateway → PaymentController', '返回支付页面/二维码', '支付系统返回支付凭证'],
        ['22', 'PaymentController → 用户', '展示支付页面', '用户完成支付操作'],
    ]
)

add_para('条件分支与异常处理：', bold=True)
add_table(
    ['场景', '触发条件', '处理方式'],
    [
        ['座位已满', '步骤5中所有座位状态为OCCUPIED', '提示"该时段已满"，返回步骤1重新选择时段'],
        ['用户时段冲突', '步骤11中发现用户已有预约', '提示"您在该时段已有预约"，返回步骤1'],
        ['座位被抢占', '步骤12中lockSeat失败', '提示"该座位已被选择"，返回步骤8重新选择'],
        ['预约创建失败', '步骤14中数据库异常', '释放座位锁定，提示"预约失败，请重试"'],
        ['支付超时', '5分钟内未完成支付', 'SchedulerService自动取消订单，释放座位，更新预约状态为EXPIRED'],
        ['支付失败', '步骤20/21中支付系统返回失败', '提示失败原因，允许重试或更换支付方式'],
        ['支付回调异常', '步骤22后回调丢失', '定时对账任务在30分钟内主动查询支付状态'],
    ]
)

# 4.3
add_title('4.3 状态图：预约状态', 2)
add_para('Reservation对象的生命周期状态变化：')
add_svg_image('5.svg', '图4-2 预约状态图', width_inches=5.0)

add_para('状态转换规则：', bold=True)
add_table(
    ['当前状态', '触发事件', '目标状态', '附加动作'],
    [
        ['PendingPayment', '支付成功回调', 'Confirmed', '确认座位分配，发送确认通知'],
        ['PendingPayment', '5分钟超时', 'Expired', '释放座位，取消订单'],
        ['PendingPayment', '用户主动取消', 'Cancelled', '释放座位，取消订单'],
        ['Confirmed', '用户扫码签到', 'InUse', '调用门禁系统开门，记录签到时间'],
        ['Confirmed', '用户取消(开始前)', 'Cancelled', '释放座位，触发退款'],
        ['Confirmed', '时段结束未签到', 'NoShow', '扣除信用分，连续3次限制预约7天'],
        ['InUse', '时段结束', 'Completed', '记录使用时长，释放座位'],
    ]
)

doc.add_page_break()

# ==================== 5. 总结 ====================
add_title('五、总结', 1)

add_para('本文档对共享自习室预约管理系统进行了完整的UML面向对象建模，包含以下内容：', indent=True)

add_para('用例建模方面：', bold=True)
add_para('识别了普通用户、管理员三个主要参与者和三个外部系统参与者，定义了11个核心用例，覆盖了预约、支付、签到、管理等全部业务场景。对"预约座位"和"预约确认与支付"两个关键用例进行了详细的事件流描述，包含基本流、备选流、异常流和业务规则。', indent=True)

add_para('静态建模方面：', bold=True)
add_para('设计了7个实体类（User、StudyRoom、Seat、TimeSlot、Reservation、Order、Violation）、12个边界类（含3个系统网关）、10个控制类和5个应用逻辑类。通过上下文类图明确了系统边界，通过完整类图展示了实体间的关系（关联、组合）和多重性，并建立了用例与类的映射关系。', indent=True)

add_para('动态建模方面：', bold=True)
add_para('以"预约座位"用例为核心绘制了通信图，展示了22条消息交互序列，涵盖座位查询、可用性校验、座位锁定、预约创建、订单生成、支付发起等完整流程。同时设计了预约状态图，定义了7种状态和7条转换规则，完整描述了预约对象的生命周期。', indent=True)

add_para('建模特点：', bold=True)
points = [
    '模型与需求一一对应，每个需求点均可在模型中找到对应元素',
    '充分考虑异常路径和边界条件，设计了完善的异常处理机制',
    '引入信用管理和定时任务等应用逻辑类，增强了系统的健壮性',
    '通过外部系统网关实现松耦合集成，符合开闭原则',
]
for pt in points:
    doc.add_paragraph(pt, style='List Bullet')

# ========== SAVE ==========
doc.save(OUTPUT)
print(f'Document saved: {OUTPUT}')
